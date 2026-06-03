#!/usr/bin/env python3

from pathlib import Path
from typing import List, Optional, Tuple

import pandas as pd

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


AGGREGATE_CSV = Path("output/benchmarks/benchmark_scaling_summary.csv")
RAW_CSV = Path("output/benchmarks/benchmark_raw_iterations_summary.csv")

COLDSTART_CSV = Path("output/benchmark_coldstart_report.csv")
WARM_BENCHMARK_CSV = Path("output/benchmark_report.csv")

OUTPUT_DOCX = Path("plots/benchmark_tables/benchmark_tables.docx")

QUERIES = ["Q1", "Q2", "Q3"]
IMPLS = ["df", "rdd"]

IMPL_LABELS = {
    "df": "DataFrame",
    "rdd": "RDD",
}

PHASE_ORDER = {
    "spark_start_s": 0,
    "loading_s": 1,
    "filtering_s": 2,
    "computation_s": 3,
    "all_airlines_computation_s": 4,
    "top10_computation_s": 5,
    "computation_percentiles_s": 6,
    "computation_minmax_s": 7,
    "output_s": 8,
    "total_s": 9,
    "wall_total_s": 10,
    "end_to_end_s": 11,
    "spark_stop_s": 12,
}

DIRECT_COMPARISON_PHASES = [
    ("spark_start_s", "Spark start"),
    ("loading_s", "Loading"),
    ("processing_s", "Processing"),
    ("wall_total_s", "End-to-end"),
    ("spark_stop_s", "Spark stop"),
]

DIRECT_COMPARISON_SOURCE_ORDER = {
    "cold-start": 0,
    "warm-session": 1,
}


def require_file(path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"Required file not found: {path}")


def normalize_aggregate_df(df: pd.DataFrame) -> pd.DataFrame:
    required = {
        "worker_count",
        "query",
        "impl",
        "phase",
        "n",
        "mean_s",
        "std_s",
        "median_s",
        "min_s",
        "max_s",
    }

    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing columns in aggregate CSV: {sorted(missing)}")

    df = df.copy()
    df["worker_count"] = pd.to_numeric(df["worker_count"], errors="coerce").astype("Int64")
    df["query"] = df["query"].astype(str).str.upper()
    df["impl"] = df["impl"].astype(str).str.lower()
    df["phase"] = df["phase"].astype(str)

    for col in ["n", "mean_s", "std_s", "median_s", "min_s", "max_s"]:
        df[col] = pd.to_numeric(df[col], errors="coerce")

    df = df.dropna(subset=["worker_count", "query", "impl", "phase", "mean_s"])
    df["worker_count"] = df["worker_count"].astype(int)

    return df


def normalize_raw_df(df: pd.DataFrame) -> pd.DataFrame:
    required = {
        "worker_count",
        "query",
        "impl",
        "iteration",
        "phase",
        "time_s",
    }

    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing columns in raw CSV: {sorted(missing)}")

    df = df.copy()
    df["worker_count"] = pd.to_numeric(df["worker_count"], errors="coerce").astype("Int64")
    df["iteration"] = pd.to_numeric(df["iteration"], errors="coerce").astype("Int64")
    df["query"] = df["query"].astype(str).str.upper()
    df["impl"] = df["impl"].astype(str).str.lower()
    df["phase"] = df["phase"].astype(str)
    df["time_s"] = pd.to_numeric(df["time_s"], errors="coerce")

    df = df.dropna(subset=["worker_count", "iteration", "query", "impl", "phase", "time_s"])
    df["worker_count"] = df["worker_count"].astype(int)
    df["iteration"] = df["iteration"].astype(int)

    return df


def normalize_phase_report_df(df: pd.DataFrame, source_name: str) -> pd.DataFrame:
    required = {
        "query",
        "phase",
        "n",
        "mean_s",
        "std_s",
        "median_s",
        "min_s",
        "max_s",
    }

    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Missing columns in {source_name}: {sorted(missing)}")

    df = df.copy()

    df["query"] = df["query"].astype(str).str.upper()
    df["phase"] = df["phase"].astype(str)

    for col in ["n", "mean_s", "std_s", "median_s", "min_s", "max_s"]:
        df[col] = pd.to_numeric(df[col], errors="coerce")

    df = df.dropna(subset=["query", "phase", "mean_s"])

    df["phase_order"] = df["phase"].map(PHASE_ORDER).fillna(999).astype(int)
    df["query_order"] = df["query"].str.extract(r"Q(\d+)")[0]
    df["query_order"] = pd.to_numeric(df["query_order"], errors="coerce").fillna(999).astype(int)

    df = df.sort_values(["query_order", "phase_order", "phase"]).reset_index(drop=True)

    return df


def format_mean_std(mean: Optional[float], std: Optional[float]) -> str:
    if mean is None or pd.isna(mean):
        return "—"

    if std is None or pd.isna(std):
        return f"{mean:.2f}"

    return f"{mean:.2f} ± {std:.2f}"


def format_value(value: Optional[float]) -> str:
    if value is None or pd.isna(value):
        return "—"
    return f"{value:.2f}"


def get_worker_counts(df: pd.DataFrame) -> List[int]:
    return sorted(int(w) for w in df["worker_count"].dropna().unique())


def aggregate_raw_phase(raw: pd.DataFrame, phase: str) -> pd.DataFrame:
    subset = raw[raw["phase"] == phase].copy()

    if subset.empty:
        raise ValueError(f"No rows found in raw CSV for phase={phase!r}")

    grouped = (
        subset.groupby(["worker_count", "query", "impl"])["time_s"]
        .agg(
            n="count",
            mean_s="mean",
            std_s="std",
            median_s="median",
            min_s="min",
            max_s="max",
        )
        .reset_index()
    )

    grouped["phase"] = phase

    return grouped[
        [
            "worker_count",
            "query",
            "impl",
            "phase",
            "n",
            "mean_s",
            "std_s",
            "median_s",
            "min_s",
            "max_s",
        ]
    ]


def build_scaling_table(
    data: pd.DataFrame,
    phase: str,
    workers: List[int],
) -> List[List[str]]:
    rows = []
    rows.append(["Query", "Impl."] + [f"{w}W" for w in workers])

    for query in QUERIES:
        for impl in IMPLS:
            row = [query, IMPL_LABELS[impl]]

            for worker in workers:
                match = data[
                    (data["worker_count"] == worker)
                    & (data["query"] == query)
                    & (data["impl"] == impl)
                    & (data["phase"] == phase)
                ]

                if match.empty:
                    row.append("—")
                else:
                    item = match.iloc[0]
                    row.append(format_mean_std(item["mean_s"], item["std_s"]))

            rows.append(row)

    return rows


def get_aggregate_mean(
    aggregate: pd.DataFrame,
    worker: int,
    query: str,
    impl: str,
    phase: str,
) -> Optional[float]:
    match = aggregate[
        (aggregate["worker_count"] == worker)
        & (aggregate["query"] == query)
        & (aggregate["impl"] == impl)
        & (aggregate["phase"] == phase)
    ]

    if match.empty:
        return None

    return float(match.iloc[0]["mean_s"])


def get_processing_mean_from_granular_aggregate(
    aggregate: pd.DataFrame,
    worker: int,
    query: str,
    impl: str,
) -> Optional[float]:
    if query == "Q1":
        return get_aggregate_mean(aggregate, worker, query, impl, "computation_s")

    if query == "Q2":
        if impl == "df":
            all_airlines = get_aggregate_mean(
                aggregate, worker, query, impl, "all_airlines_computation_s"
            )
            top10 = get_aggregate_mean(
                aggregate, worker, query, impl, "top10_computation_s"
            )

            if all_airlines is None or top10 is None:
                return None

            return all_airlines + top10

        if impl == "rdd":
            return get_aggregate_mean(aggregate, worker, query, impl, "computation_s")

    if query == "Q3":
        return get_aggregate_mean(
            aggregate, worker, query, impl, "computation_percentiles_s"
        )

    return None


def build_phase_breakdown_table(
    aggregate: pd.DataFrame,
    worker: int,
) -> List[List[str]]:
    rows = []
    rows.append(
        [
            "Query",
            "Impl.",
            "Loading",
            "Filtering",
            "Processing",
            "Output",
            "End-to-end",
        ]
    )

    for query in QUERIES:
        for impl in IMPLS:
            loading = get_aggregate_mean(aggregate, worker, query, impl, "loading_s")
            filtering = get_aggregate_mean(aggregate, worker, query, impl, "filtering_s")
            processing = get_processing_mean_from_granular_aggregate(
                aggregate, worker, query, impl
            )
            output = get_aggregate_mean(aggregate, worker, query, impl, "output_s")
            end_to_end = get_aggregate_mean(aggregate, worker, query, impl, "end_to_end_s")

            rows.append(
                [
                    query,
                    IMPL_LABELS[impl],
                    format_value(loading),
                    format_value(filtering),
                    format_value(processing),
                    format_value(output),
                    format_value(end_to_end),
                ]
            )

    return rows


def format_phase_name(phase: str) -> str:
    labels = {
        "spark_start_s": "Spark start",
        "loading_s": "Loading",
        "filtering_s": "Filtering",
        "computation_s": "Computation",
        "all_airlines_computation_s": "All-airlines computation",
        "top10_computation_s": "Top-10 computation",
        "computation_percentiles_s": "Percentiles computation",
        "computation_minmax_s": "Min/max computation",
        "output_s": "Output",
        "total_s": "Total",
        "wall_total_s": "Wall total",
        "end_to_end_s": "End-to-end",
        "spark_stop_s": "Spark stop",
    }

    return labels.get(phase, phase)


def build_phase_summary_rows(df: pd.DataFrame) -> List[List[str]]:
    rows = [
        [
            "Query",
            "Phase",
            "n",
            "Mean (s)",
            "StdDev",
            "Median",
            "Min",
            "Max",
        ]
    ]

    for _, row in df.iterrows():
        rows.append(
            [
                row["query"],
                format_phase_name(row["phase"]),
                str(int(row["n"])) if not pd.isna(row["n"]) else "—",
                format_value(row["mean_s"]),
                format_value(row["std_s"]),
                format_value(row["median_s"]),
                format_value(row["min_s"]),
                format_value(row["max_s"]),
            ]
        )

    return rows


def get_phase_row(df: pd.DataFrame, query: str, phase: str) -> Optional[pd.Series]:
    match = df[
        (df["query"] == query)
        & (df["phase"] == phase)
    ]

    if match.empty:
        return None

    return match.iloc[0]


def sum_phase_rows(
    first: pd.Series,
    second: pd.Series,
    query: str,
    phase: str,
) -> dict:
    """
    Sum two aggregate rows element by element.

    This is intentionally element-wise because the source files only contain
    aggregate statistics, not per-iteration raw timings.
    """

    n_first = first.get("n")
    n_second = second.get("n")

    if pd.isna(n_first):
        n = n_second
    elif pd.isna(n_second):
        n = n_first
    else:
        n = min(int(n_first), int(n_second))

    return {
        "query": query,
        "phase": phase,
        "n": n,
        "mean_s": first["mean_s"] + second["mean_s"],
        "std_s": first["std_s"] + second["std_s"],
        "median_s": first["median_s"] + second["median_s"],
        "min_s": first["min_s"] + second["min_s"],
        "max_s": first["max_s"] + second["max_s"],
    }


def get_normalized_report_phase_row(
    df: pd.DataFrame,
    query: str,
    target_phase: str,
) -> Optional[dict]:
    """
    Return one normalized row for the direct cold-start vs warm-session table.

    Kept logical phases:
    - spark_start_s
    - loading_s
    - processing_s
    - wall_total_s
    - spark_stop_s

    Q2 processing is computed as:
    all_airlines_computation_s + top10_computation_s
    element by element.
    """

    if target_phase == "processing_s":
        if query == "Q2":
            all_airlines = get_phase_row(df, query, "all_airlines_computation_s")
            top10 = get_phase_row(df, query, "top10_computation_s")

            if all_airlines is not None and top10 is not None:
                return sum_phase_rows(
                    first=all_airlines,
                    second=top10,
                    query=query,
                    phase="processing_s",
                )

            fallback = get_phase_row(df, query, "computation_s")
            if fallback is None:
                return None

            return {
                "query": query,
                "phase": "processing_s",
                "n": fallback["n"],
                "mean_s": fallback["mean_s"],
                "std_s": fallback["std_s"],
                "median_s": fallback["median_s"],
                "min_s": fallback["min_s"],
                "max_s": fallback["max_s"],
            }

        if query == "Q3":
            row = get_phase_row(df, query, "computation_s")

            if row is None:
                row = get_phase_row(df, query, "computation_percentiles_s")

            if row is None:
                return None

            return {
                "query": query,
                "phase": "processing_s",
                "n": row["n"],
                "mean_s": row["mean_s"],
                "std_s": row["std_s"],
                "median_s": row["median_s"],
                "min_s": row["min_s"],
                "max_s": row["max_s"],
            }

        row = get_phase_row(df, query, "computation_s")

        if row is None:
            return None

        return {
            "query": query,
            "phase": "processing_s",
            "n": row["n"],
            "mean_s": row["mean_s"],
            "std_s": row["std_s"],
            "median_s": row["median_s"],
            "min_s": row["min_s"],
            "max_s": row["max_s"],
        }

    row = get_phase_row(df, query, target_phase)

    if row is None:
        return None

    return {
        "query": query,
        "phase": target_phase,
        "n": row["n"],
        "mean_s": row["mean_s"],
        "std_s": row["std_s"],
        "median_s": row["median_s"],
        "min_s": row["min_s"],
        "max_s": row["max_s"],
    }

def build_cold_vs_warm_rows(
    coldstart: pd.DataFrame,
    warm: pd.DataFrame,
) -> Tuple[List[List[str]], List[Tuple[str, str]]]:
    """
    Build a direct comparison table between cold-start and warm-session reports.

    Visible columns:
    query, phase, mode, mean_s, std_s, median_s, min_s, max_s
    """

    rows: List[List[str]] = [
        ["Query", "Phase", "Mode", "Mean (s)", "StdDev", "Median", "Min", "Max"]
    ]

    group_keys: List[Tuple[str, str]] = [("__header__", "__header__")]

    sources = [
        ("cold start", coldstart),
        ("warmup", warm),
    ]

    for query in QUERIES:
        for target_phase, phase_label in DIRECT_COMPARISON_PHASES:
            for mode_label, source_df in sources:
                normalized = get_normalized_report_phase_row(
                    df=source_df,
                    query=query,
                    target_phase=target_phase,
                )

                if normalized is None:
                    continue

                rows.append(
                    [
                        query,
                        phase_label,
                        mode_label,
                        format_value(normalized["mean_s"]),
                        format_value(normalized["std_s"]),
                        format_value(normalized["median_s"]),
                        format_value(normalized["min_s"]),
                        format_value(normalized["max_s"]),
                    ]
                )

                group_keys.append((query, phase_label))

    return rows, group_keys


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")

    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge_name, edge_data in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        if edge_data is None:
            continue

        tag = "w:" + edge_name
        element = tc_borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)

        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_cell_borders(cell) -> None:
    no_border = {"val": "nil", "sz": "0", "color": "FFFFFF"}
    set_cell_borders(
        cell,
        top=no_border,
        bottom=no_border,
        left=no_border,
        right=no_border,
    )


def set_table_academic_borders(table) -> None:
    strong = {"val": "single", "sz": "12", "color": "000000"}
    light = {"val": "single", "sz": "4", "color": "BFBFBF"}
    none = {"val": "nil", "sz": "0", "color": "FFFFFF"}

    row_count = len(table.rows)

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            clear_cell_borders(cell)

            top = None
            bottom = None

            if r_idx == 0:
                top = strong
                bottom = strong
            elif r_idx == row_count - 1:
                bottom = strong
            else:
                bottom = light

            set_cell_borders(
                cell,
                top=top,
                bottom=bottom,
                left=none,
                right=none,
            )


def set_table_academic_borders_with_groups(table, rows: List[List[str]]) -> None:
    strong = {"val": "single", "sz": "12", "color": "000000"}
    medium = {"val": "single", "sz": "10", "color": "000000"}
    light = {"val": "single", "sz": "4", "color": "BFBFBF"}
    none = {"val": "nil", "sz": "0", "color": "FFFFFF"}

    row_count = len(table.rows)

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            clear_cell_borders(cell)

            top = None
            bottom = light

            if r_idx == 0:
                top = strong
                bottom = strong
            elif r_idx == row_count - 1:
                bottom = strong
            else:
                current_group = rows[r_idx][0]
                previous_group = rows[r_idx - 1][0]

                if current_group != previous_group:
                    top = medium

            set_cell_borders(
                cell,
                top=top,
                bottom=bottom,
                left=none,
                right=none,
            )


def set_table_borders_by_query_and_phase(
    table,
    group_keys: List[Tuple[str, str]],
) -> None:
    """
    Stronger border between queries.
    Medium border between phases inside the same query.
    No special border between cold-start and warm-session rows of the same phase.
    """

    strong = {"val": "single", "sz": "14", "color": "000000"}
    medium = {"val": "single", "sz": "10", "color": "000000"}
    light = {"val": "single", "sz": "4", "color": "BFBFBF"}
    none = {"val": "nil", "sz": "0", "color": "FFFFFF"}

    row_count = len(table.rows)

    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            clear_cell_borders(cell)

            top = None
            bottom = light

            if r_idx == 0:
                top = strong
                bottom = strong
            elif r_idx == row_count - 1:
                bottom = strong
            else:
                current_query, current_phase = group_keys[r_idx]
                previous_query, previous_phase = group_keys[r_idx - 1]

                if current_query != previous_query:
                    top = strong
                elif current_phase != previous_phase:
                    top = medium

            set_cell_borders(
                cell,
                top=top,
                bottom=bottom,
                left=none,
                right=none,
            )


def set_cell_text(cell, text: str, bold: bool = False, align_center: bool = True) -> None:
    cell.text = ""

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_academic_table(
    document: Document,
    caption: str,
    rows: List[List[str]],
    group_by_first_column: bool = False,
    group_keys: Optional[List[Tuple[str, str]]] = None,
) -> None:
    add_caption(document, caption)

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]

        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            is_header = r_idx == 0
            set_cell_text(cell, value, bold=is_header, align_center=True)

            if is_header:
                set_cell_shading(cell, "F2F2F2")

    if group_keys is not None:
        set_table_borders_by_query_and_phase(table, group_keys)
    elif group_by_first_column:
        set_table_academic_borders_with_groups(table, rows)
    else:
        set_table_academic_borders(table)

    document.add_paragraph()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)


def main() -> None:
    require_file(AGGREGATE_CSV)
    require_file(RAW_CSV)
    require_file(COLDSTART_CSV)
    require_file(WARM_BENCHMARK_CSV)

    aggregate = normalize_aggregate_df(pd.read_csv(AGGREGATE_CSV))
    raw = normalize_raw_df(pd.read_csv(RAW_CSV))

    coldstart = normalize_phase_report_df(
        pd.read_csv(COLDSTART_CSV),
        source_name=str(COLDSTART_CSV),
    )

    warm_benchmark = normalize_phase_report_df(
        pd.read_csv(WARM_BENCHMARK_CSV),
        source_name=str(WARM_BENCHMARK_CSV),
    )

    workers = get_worker_counts(aggregate)

    if not workers:
        raise ValueError("No worker_count values found in aggregate CSV.")

    if 1 not in workers:
        raise ValueError("worker_count=1 not found in aggregate CSV; cannot build Table 3.")

    end_to_end_rows = build_scaling_table(
        data=aggregate,
        phase="end_to_end_s",
        workers=workers,
    )

    processing_aggregate = aggregate_raw_phase(raw, phase="processing_s")

    processing_rows = build_scaling_table(
        data=processing_aggregate,
        phase="processing_s",
        workers=workers,
    )

    phase_breakdown_rows = build_phase_breakdown_table(
        aggregate=aggregate,
        worker=1,
    )

    coldstart_rows = build_phase_summary_rows(coldstart)
    warm_benchmark_rows = build_phase_summary_rows(warm_benchmark)

    cold_vs_warm_rows, cold_vs_warm_group_keys = build_cold_vs_warm_rows(
        coldstart=coldstart,
        warm=warm_benchmark,
    )

    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Benchmark Summary Tables")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)

    document.add_paragraph()

    add_academic_table(
        document,
        "Table 1: End-to-end execution time scaling results",
        end_to_end_rows,
    )

    add_note(
        document,
        "Values are reported as mean ± standard deviation over the valid benchmark iterations.",
    )

    document.add_paragraph()

    add_academic_table(
        document,
        "Table 2: Processing time scaling results",
        processing_rows,
    )

    add_note(
        document,
        "Processing time is computed from the normalized processing_s phase in the raw benchmark summary. "
        "For Q2 DataFrame, processing_s includes both all-airlines aggregation and top-10 extraction.",
    )

    document.add_paragraph()

    add_academic_table(
        document,
        "Table 3: Average phase breakdown with 1 Spark worker",
        phase_breakdown_rows,
    )

    add_note(
        document,
        "Values are mean execution times in seconds. Processing is normalized per query: "
        "Q1 uses computation_s, Q2 DataFrame uses all_airlines_computation_s + top10_computation_s, "
        "Q2 RDD uses computation_s, and Q3 uses computation_percentiles_s.",
    )

    document.add_paragraph()

    add_academic_table(
        document,
        "Table 4: Cold-start benchmark phase summary",
        coldstart_rows,
        group_by_first_column=True,
    )

    add_note(
        document,
        "Values are mean execution times in seconds. The cold-start benchmark includes "
        "Spark session startup and shutdown phases when available.",
    )

    document.add_paragraph()

    add_academic_table(
        document,
        "Table 5: Warm-session benchmark phase summary",
        warm_benchmark_rows,
        group_by_first_column=True,
    )

    add_note(
        document,
        "Values are mean execution times in seconds. The warm-session benchmark reuses "
        "the same Spark session across iterations and reports the measured phases for each query.",
    )

    document.add_paragraph()

    add_academic_table(
        document,
        "Table 6: Cold-start vs warm-session benchmark phase comparison",
        cold_vs_warm_rows,
        group_keys=cold_vs_warm_group_keys,
    )

    add_note(
        document,
        "Only selected phases are reported: Spark start, loading, processing, end-to-end and Spark stop. "
        "For Q2, processing is computed as all_airlines_computation_s + top10_computation_s, summed element by element "
        "from the aggregate rows. The mode column distinguishes cold start from warmup execution.",
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)

    print(f"[OK] Saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()